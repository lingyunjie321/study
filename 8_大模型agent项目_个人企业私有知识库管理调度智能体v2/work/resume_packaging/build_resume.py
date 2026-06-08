from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
OUT_DOCX = OUT_DIR / "凌云杰_AI应用开发简历_企业知识问答优化版.docx"

FONT_CN = "Microsoft YaHei"
FONT_LATIN = "Calibri"
BLUE = RGBColor(31, 77, 120)
LIGHT_BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(90, 90, 90)
INK = RGBColor(30, 30, 30)


def set_run_font(run, size=None, bold=None, color=None, italic=None, name=FONT_CN):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def paragraph_border_bottom(paragraph, color="D9E2F3", size="8", space="3"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)


def set_para_format(paragraph, before=0, after=2, line=1.05):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_text(paragraph, text, size=9.2, bold=False, color=INK, italic=False):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, italic=italic)
    return run


def add_heading(doc, text):
    p = doc.add_paragraph()
    set_para_format(p, before=8, after=3, line=1.0)
    paragraph_border_bottom(p)
    add_text(p, text, size=11.5, bold=True, color=BLUE)
    return p


def add_line(doc, text="", size=9.2, bold=False, color=INK, after=2, before=0, align=None):
    p = doc.add_paragraph()
    set_para_format(p, before=before, after=after)
    if align is not None:
        p.alignment = align
    add_text(p, text, size=size, bold=bold, color=color)
    return p


def add_rich_line(doc, parts, after=2, before=0, align=None):
    p = doc.add_paragraph()
    set_para_format(p, before=before, after=after)
    if align is not None:
        p.alignment = align
    for part in parts:
        add_text(
            p,
            part["text"],
            size=part.get("size", 9.2),
            bold=part.get("bold", False),
            color=part.get("color", INK),
            italic=part.get("italic", False),
        )
    return p


def add_bullet(doc, text, size=9.0, after=1.4):
    p = doc.add_paragraph(style="List Bullet")
    set_para_format(p, before=0, after=after, line=1.04)
    fmt = p.paragraph_format
    fmt.left_indent = Inches(0.21)
    fmt.first_line_indent = Inches(-0.13)
    add_text(p, text, size=size, color=INK)
    return p


def add_project(doc, title, meta, direction, stack, bullets, results=None):
    add_rich_line(
        doc,
        [
            {"text": title, "size": 9.8, "bold": True, "color": BLUE},
            {"text": f"  {meta}", "size": 8.8, "bold": True, "color": MUTED},
        ],
        before=3,
        after=1,
    )
    add_rich_line(
        doc,
        [
            {"text": "项目方向：", "size": 8.8, "bold": True, "color": MUTED},
            {"text": direction, "size": 8.8, "color": MUTED},
        ],
        after=0.6,
    )
    add_rich_line(
        doc,
        [
            {"text": "技术关键词：", "size": 8.8, "bold": True, "color": MUTED},
            {"text": stack, "size": 8.8, "color": MUTED},
        ],
        after=1.2,
    )
    for item in bullets:
        add_bullet(doc, item)
    if results:
        add_rich_line(doc, [{"text": "成果：", "size": 9.0, "bold": True, "color": BLUE}], after=0.3)
        for item in results:
            add_bullet(doc, item, size=8.9)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_CN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.05

    for style_name in ["List Bullet"]:
        st = styles[style_name]
        st.font.name = FONT_CN
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        st.font.size = Pt(9.0)
        st.paragraph_format.space_after = Pt(1.5)
        st.paragraph_format.line_spacing = 1.04

    # Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_format(p, before=0, after=1, line=1.0)
    add_text(p, "凌云杰", size=18, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_format(p, before=0, after=2, line=1.0)
    add_text(p, "AI 应用开发 / 算法工程师", size=10.2, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_format(p, before=0, after=4, line=1.0)
    add_text(p, "男 | 13277855173 | yunleing1227@163.com | 5年工作经验 | 深圳 | 期望薪资：20-40K", size=8.8, color=MUTED)
    paragraph_border_bottom(p, color="B4C7E7", size="10", space="4")

    add_heading(doc, "个人概要")
    add_bullet(doc, "具备企业级 AI 应用开发、Agent 工作流、RAG 知识检索、NL2SQL、MCP 工具接入与计算机视觉落地经验，能够从业务需求、数据上下文、工程实现到服务化交付完成闭环。")
    add_bullet(doc, "AI 应用经历从地产公司内部知识问答原型开始，后续在半导体企业场景中持续参与数据工程 Agent、Agentic RL 训练评估平台、芯片图像自动量测工具建设。")
    add_bullet(doc, "关注企业真实使用中的可控性问题，包括知识来源引用、业务口径约束、工具权限、安全执行、结果校验、日志追踪和多入口交付。")

    add_heading(doc, "核心技能")
    skill_lines = [
        ("AI 应用 / Agent", "RAG、GraphRAG、Agent Workflow、Tool Calling、Prompt Engineering、MCP、Skills、权限治理、自动化评测"),
        ("数据智能", "NL2SQL、Schema Linking、SQLGlot、参考 SQL 注入、指标口径检索、企业数据分析助手"),
        ("工程开发", "Python、FastAPI、Pydantic、REST API、CLI、Docker、Web/TUI、服务化部署、单元测试"),
        ("存储与检索", "LanceDB、ChromaDB/FAISS、Neo4j、向量检索、文档分块、元数据过滤、增量更新"),
        ("模型与视觉", "PyTorch、verl、Ray、FSDP、vLLM、PPO/GRPO、MMSegmentation、OpenCV、语义分割、图像后处理"),
    ]
    for label, value in skill_lines:
        add_rich_line(
            doc,
            [
                {"text": f"{label}：", "size": 9.0, "bold": True, "color": BLUE},
                {"text": value, "size": 9.0, "color": INK},
            ],
            after=1.2,
        )

    add_heading(doc, "工作经历")
    add_rich_line(
        doc,
        [
            {"text": "海思半导体科技有限公司（派遣）", "size": 9.8, "bold": True, "color": BLUE},
            {"text": "  算法工程师  2024.04-2026.06", "size": 9.0, "bold": True, "color": MUTED},
        ],
        before=2,
        after=1,
    )
    hisi_bullets = [
        "从事企业级 AI 应用开发与 AI Agent 平台研发，服务于内部数据分析、知识检索、自然语言取数、自动化工具调用及芯片图像量测等场景。",
        "参与企业数据工程 Agent 平台建设，将自然语言问题拆解为 Schema 检索、上下文注入、SQL 生成、执行校验和结果输出等环节，降低表名幻觉、字段误用和业务口径偏差。",
        "参与 Agentic RL 训练与评估平台建设，围绕多工具 Rollout、Reward 设计、模型部署和自动化评测，探索提升 Agent 多步推理、工具使用和结果校验能力。",
        "参与半导体芯片图像自动量测项目，负责语义分割模型训练、图像后处理和量测算子开发，在部分规则清晰场景中实现约 15 分钟完成 50 张图像批量量测。",
    ]
    for item in hisi_bullets:
        add_bullet(doc, item)

    add_rich_line(
        doc,
        [
            {"text": "广西南宁冠能投资有限公司", "size": 9.8, "bold": True, "color": BLUE},
            {"text": "  IT 技术支持  2021.02-2024.01", "size": 9.0, "bold": True, "color": MUTED},
        ],
        before=4,
        after=1,
    )
    guanneng_bullets = [
        "负责 OA 办公系统日常维护、功能二次开发、问题测试与故障响应，保障营销、工程、物业、财务等部门的信息化系统稳定运行。",
        "管理公司 IT 工具、PC 设备和电子数据，搭建 IT 报修工单流程，响应全公司软硬件及系统问题咨询。",
        "长期对接业务部门信息化需求，参与处理销控房源数据不一致、工程款支付异常、业主缴费系统故障等跨部门业务卡点。",
        "2023 年下半年参与地产公司内部知识问答系统原型建设，主要承担业务资料梳理、知识入库、问题集建设、问答效果验证和部分接口联调，该项目成为本人转向 AI 应用开发的启蒙项目。",
    ]
    for item in guanneng_bullets:
        add_bullet(doc, item)

    add_heading(doc, "项目经历")
    add_project(
        doc,
        "可插拔 Agentic RL 训练与评估平台",
        "AI应用开发算法工程师 | 2025.10-2026.06",
        "Agent 强化学习、工具调用训练、训练评估闭环、行业数据构建",
        "PyTorch、verl、Ray、FSDP、vLLM、PPO/GRPO、SFT、Tool Calling、Reward",
        [
            "搭建 Agentic RL 训练链路，覆盖数据加载、SFT 冷启动、Rollout 生成、工具调用、Reward 计算、策略更新、Checkpoint 管理、模型转换和自动化评估。",
            "实现多工具交互式 Rollout 能力，支持模型在推理过程中调用搜索工具和 Python 计算工具，并将工具结果回填上下文继续生成。",
            "设计可插拔工具与任务配置接口，将搜索、计算、评估、数据处理等能力抽象为可配置组件，支持按任务替换工具、数据集和奖励规则。",
            "参与 PPO/GRPO 策略优化逻辑改造，实现可配置的 entropy-aware advantage、adaptive clipping 和动态 Rollout 分支采样逻辑。",
        ],
        [
            "打通 SFT、Agentic RL、多工具 Rollout、自动化评估的端到端流程，为后续迁移到半导体良率分析、设备异常排查、供应链风险分析等场景提供基础。",
        ],
    )

    add_project(
        doc,
        "企业级数据工程 Agent 与 AI Harness 平台",
        "Python | 2025.01-2026.06",
        "企业数据智能、NL2SQL、RAG、Agent Runtime、MCP、工具治理",
        "Python、FastAPI、LanceDB、SQLGlot、MCP、Skills、Agent Workflow、Tool Permission",
        [
            "设计并实现基于 Workflow/Node 的 Agent 执行链路，将自然语言取数拆解为任务初始化、Schema 检索、上下文注入、SQL 生成、执行验证和结果输出。",
            "建设多源数据上下文体系，将表结构、字段说明、业务指标、参考 SQL、业务文档和外部知识统一纳入 RAG 存储，为 SQL 生成和结果解释提供可检索上下文。",
            "参与 GenSQL Agentic Node 设计，支持数据库工具、上下文检索工具、参考 SQL 工具、日期解析工具、文件工具和子 Agent 任务工具按配置加载。",
            "实现 MCP 工具生态接入和 Skills 插件化能力，将数据查询、知识检索、SQL 分析、报表分析等能力封装为可发现、可加载、可授权的技能模块。",
            "参与工具权限治理与执行安全机制，对数据库查询、文件访问、MCP 工具、技能加载和脚本调用进行白名单、人工确认、会话级授权、超时控制和敏感路径隔离。",
        ],
        [
            "打通“自然语言问题 -> 数据上下文检索 -> SQL 生成 -> 执行验证 -> 结果输出”的企业数据分析 Agent 链路。",
            "将 Agent 能力从单一问答扩展为可配置 Runtime，支持工具加载、MCP 接入、Skills 复用、权限确认、执行日志和多入口服务化交付。",
        ],
    )

    add_project(
        doc,
        "半导体芯片图像自动量测 AI 工具",
        "算法工程师 | 2024.05-2026.06",
        "半导体芯片量测自动化、语义分割、图像后处理、批量量测",
        "Python、PyTorch、MMSegmentation、OpenCV、语义分割、模型训练、模型推理、量测算子",
        [
            "基于业务方量测需求梳理芯片图像场景，明确目标结构、分割类别、量测点位、长度计算口径和结果输入输出格式。",
            "参与芯片图像数据集建设，包括场景划分、样本筛选、标注规范制定和标注质量检查，覆盖目标结构边界、噪声、灰度差异和典型工艺形态。",
            "基于 MMSegmentation 完成特定场景语义分割模型训练，参与训练配置、数据增强、推理效果验证和误差样本分析。",
            "开发图像后处理与量测算子，根据分割图提取目标区域、边界、中心线或关键点位，结合业务规则计算指定点位长度并转换为业务可读结果。",
            "建设模型推理与后端处理流程，支持前端上传指定场景图像后自动选择对应模型完成推理、分割图生成、量测算子执行和结果返回。",
        ],
        [
            "打通“芯片图像上传 -> 场景模型推理 -> 分割图生成 -> 量测算子计算 -> 量测结果返回”的端到端自动化量测链路。",
            "在量测规则相对清晰、图像结构不复杂的场景中，实现约 15 分钟完成 50 张图像批量量测，显著优于传统人工逐张量测流程。",
        ],
    )

    add_project(
        doc,
        "地产公司企业内部知识问答系统原型",
        "IT技术支持 / AI应用开发参与 | 2023.08-2024.01",
        "企业知识库问答、RAG 原型、内部制度/流程检索、AI 应用启蒙项目",
        "Python、FastAPI、LangChain、ChromaDB/FAISS、OpenAI 兼容模型、Embedding、文档解析、Prompt、HTML/JS",
        [
            "参与营销、工程、物业、财务、OA 等部门知识资料梳理，将制度文档、流程说明、合同模板、常见问题、Excel 台账等资料按部门、主题、版本和来源进行分类。",
            "协助建设文档入库流程，参与 PDF/Word/Excel/TXT 等资料的解析、文本清洗、分块、元数据标注和向量化入库，使内部资料具备可检索基础。",
            "参与 RAG 问答链路联调，围绕用户问题改写、向量召回、上下文拼接、Prompt 约束、答案来源引用和“无答案时不编造”等策略进行测试与反馈。",
            "配合业务方建设高频问题集，覆盖销控房源、工程款流程、物业缴费、OA 审批、合同模板查询等场景，用于验证问答效果和记录失败样例。",
            "参与简单 Web/API 原型联调，支持资料上传、问题输入、答案返回和来源展示，帮助业务人员以自然语言查询内部制度和流程资料。",
        ],
        [
            "打通地产公司内部知识资料从“人工问询/文件夹查找”到“自然语言检索 + 来源可追溯回答”的原型链路。",
            "项目以参与和学习为主，帮助本人建立对 RAG、Embedding、向量检索、Prompt 约束、文档入库和 AI 应用工程化的第一层实践认知，并成为后续转向 AI 应用开发的起点。",
        ],
    )

    add_heading(doc, "个人优势")
    strengths = [
        "AI 应用落地能力：不局限于大模型 API 调用，能够围绕企业真实业务完成 RAG、工具调用、工作流编排、服务化接入、权限治理和效果评估。",
        "Agent 工程化能力：具备 Agent Runtime、MCP、Skills、工具权限、会话状态、多入口交付等实践经验，关注 Agent 的可控性和可复用性。",
        "垂直行业迁移能力：具备地产、半导体等行业场景经验，能够将知识问答、数据分析、图像识别和自动化执行能力迁移到具体业务流程。",
        "从业务到技术的沟通能力：早期 IT 支持经历使本人熟悉跨部门需求梳理、问题定位、用户反馈收集和系统落地推进。",
    ]
    for item in strengths:
        add_bullet(doc, item)

    add_heading(doc, "教育经历 / 证书")
    add_rich_line(
        doc,
        [
            {"text": "广西大学行健文理学院", "size": 9.4, "bold": True, "color": BLUE},
            {"text": "  本科  土木工程  2016-2020", "size": 9.0, "color": INK},
        ],
        after=1,
    )
    add_line(doc, "大学英语四级", size=9.0, after=0)

    # Footer
    for sec in doc.sections:
        footer_p = sec.footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_format(footer_p, before=0, after=0, line=1.0)
        add_text(footer_p, "凌云杰 | AI 应用开发 / 算法工程师", size=8, color=MUTED)

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
